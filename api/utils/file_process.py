from docx import Document
from docx.text.paragraph import Paragraph
import os
import json

from config import base_path

# Define mappings for different levels of work units
level1 = level2 = {"ASE_INT.1-1": {}, "ASE_INT.1-2": {}, "ASE_INT.1-3": {}, "ASE_INT.1-4": {}, "ASE_INT.1-5": {},
                   "ASE_INT.1-6": {}, "ASE_INT.1-7": {}, "ASE_INT.1-8": {}, "ASE_INT.1-9": {}, "ASE_INT.1-10": {},
                   "ASE_INT.1-11": {}, "ASE_OBJ.1-1": {}, "ASE_REQ.3-1": {}, "ASE_REQ.3-2": {}, "ASE_REQ.3-3": {},
                   "ASE_REQ.3-4": {}, "ASE_REQ.3-5": {}, "ASE_REQ.3-6": {}, "ASE_REQ.3-7": {}, "ASE_TSS.1-1": {},
                   "ASE_TSS.1-2": {}, "ALC_FLR.2-1": {}, "ALC_FLR.2-2": {}, "ALC_FLR.2-3": {}, "ALC_FLR.2-4": {},
                   "ALC_FLR.2-5": {}, "ALC_FLR.2-6": {}, "ALC_FLR.2-7": {}, "ALC_FLR.2-8": {}, "ALC_FLR.2-9": {},
                   "ALC_FLR.2-10": {}}

level_unit_mapping = {1: level1, 2: level2}


def validate_unit_name(unit_name: str, mapping: dict):
    """
    Validate if the given work unit name exists in the provided mapping.

    Args:
        unit_name (str): The name of the work unit to validate.
        mapping (dict): A dictionary containing valid work unit names as keys.

    Returns:
        bool: True if the unit name exists in the mapping, False otherwise.
    """

    return unit_name in mapping.keys()


def get_files(dir_path: str, sesip_level: int):
    """
    Load the report template and evaluation details based on the security target ID and SESIP level.

    Args:
        dir_path (str): The directory path of the evaluation details file.
        sesip_level (int): The SESIP level used to locate the report template.

    Returns:
        tuple: A tuple containing the loaded Document object for the report template and
               a dictionary of evaluation details.
    """

    template_path = os.path.join(base_path, f"evaluation_report_template_level_{sesip_level}.docx")
    details_path = os.path.join(dir_path, "eval_details.json")

    report_template = Document(template_path)
    eval_details = json.load(open(details_path, "r"))

    return report_template, eval_details


def map_units(work_units: list, mapping: dict):
    """
    Map the work units from the evaluation details to the provided mapping.

    Args:
        work_units (list): A list of work unit dictionaries from the evaluation details.
        mapping (dict): A dictionary to map work unit names to their descriptions and statuses.

    Modifies:
        mapping (dict): Updates the mapping dictionary with descriptions and statuses of work units.
    """

    for unit in work_units:
        unit_name = unit["Work_Unit_Name"]
        if validate_unit_name(unit_name, mapping):
            mapping[unit_name] = {
                "description": unit["Work_Unit_Description"],
                "status": unit["Work_Unit_Evaluation_Result_Status"]
            }


def update_status(p: Paragraph, status: str):
    """
    Update the status text in a paragraph, setting it to bold and applying the 'Times New Roman' font.

    Args:
        p (Paragraph): The paragraph to update.
        status (str): The status text to add to the paragraph.

    Modifies:
        p (Paragraph): Updates the paragraph by adding the status text in bold and with the specified font.
    """

    status = p.add_run(status)
    status.bold = True
    status.font.name = "Times New Roman"


def update_description(p: Paragraph, description: str):
    """
    Update the description text in a paragraph, applying the 'Times New Roman' font.

    Args:
        p (Paragraph): The paragraph to update.
        description (str): The description text to add to the paragraph.

    Modifies:
        p (Paragraph): Updates the paragraph by adding the description text with the specified font.
    """

    description = p.add_run(description)
    description.font.name = "Times New Roman"


def update_document(report_template: Document, unit_mapping: dict):
    """
    Update the document with the work unit status and descriptions based on the provided mapping.

    Args:
        report_template (Document): The Document object representing the report template.
        unit_mapping (dict): A dictionary mapping work unit names to their statuses and descriptions.

    Modifies:
        report_template (Document): Updates the report template with status and description information.
    """

    p = report_template.paragraphs
    not_found = []

    for i in range(len(p)):
        unit_name = p[i].text.split("(")[0].strip()
        if validate_unit_name(unit_name, unit_mapping):
            try:
                update_status(p[i + 2], unit_mapping[unit_name]["status"])
                update_description(p[i + 3], unit_mapping[unit_name]["description"])

            except KeyError:
                not_found.append(unit_name)
                pass

    if not_found:
        print(f"Could not find: {', '.join(not_found)}")


def generate_eval_report(dir_path: str, sesip_level: int):
    """
    Generate an evaluation report by loading the template, mapping work units, and updating the document.

    Args:
        dir_path (str): The directory path of the evaluation details file.
        sesip_level (int): The SESIP level used to select the appropriate report template.
    """

    report_template, eval_details = get_files(st_id, sesip_level)

    work_units = eval_details["Work_Units"]
    unit_mapping = level_unit_mapping[sesip_level]
    map_units(work_units, unit_mapping)

    update_document(report_template, unit_mapping)
    report_template.save(os.path.join(dir_path, "eval_file.docx"))
